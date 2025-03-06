from docx import Document
from typing import Dict, Any, List
from ..core.logger import log_info, log_error

class DocxProcessor:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def extract_data(self) -> Dict[str, Any]:
        """Extract text, tables and metadata from DOCX file."""
        try:
            doc = Document(self.file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Only add rows that aren't completely empty
                        table_data.append(row_data)
                if table_data:  # Only add tables that aren't completely empty
                    tables.append(table_data)

            # Extract document properties
            properties = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'last_modified_by': doc.core_properties.last_modified_by or ''
            }

            # Extract sections and their properties
            sections = []
            for section in doc.sections:
                sections.append({
                    'start_type': str(section.start_type),
                    'orientation': str(section.orientation),
                    'page_height': section.page_height.inches,
                    'page_width': section.page_width.inches
                })

            return {
                'text': '\n'.join(paragraphs),
                'tables': tables,
                'metadata': {
                    'properties': properties,
                    'sections': sections,
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables)
                }
            }

        except Exception as e:
            log_error(f"Error processing DOCX {self.file_path}: {str(e)}")
            return {
                'text': '',
                'tables': [],
                'metadata': {}
            }

    def extract_text_only(self) -> str:
        """Extract only text content from the document."""
        try:
            doc = Document(self.file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        except Exception as e:
            log_error(f"Error extracting text from DOCX {self.file_path}: {str(e)}")
            return ''